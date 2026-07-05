from __future__ import annotations

import json
import shutil
import sys
import zipfile
from pathlib import Path

from docx import Document


def paragraph_text(paragraph) -> str:
    return clean_text("".join(run.text for run in paragraph.runs))


def clean_text(value: str) -> str:
    text = value.strip()
    if not text:
        return ""
    try:
        repaired = text.encode("latin1").decode("utf-8")
    except UnicodeError:
        repaired = text
    return " ".join(repaired.replace("\u00a0", " ").split())


def iter_block_items(document: Document):
    body = document.element.body
    paragraphs = {paragraph._p: paragraph for paragraph in document.paragraphs}
    tables = {table._tbl: table for table in document.tables}

    for child in body.iterchildren():
        if child in paragraphs:
            text = paragraph_text(paragraphs[child])
            if text:
                yield {"type": "paragraph", "text": text}
        elif child in tables:
            rows = []
            for row in tables[child].rows:
                rows.append([clean_text(cell.text) for cell in row.cells])
            yield {"type": "table", "rows": rows}


def extract_images(docx_path: Path, output_dir: Path) -> list[dict[str, str]]:
    images_dir = output_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    images = []

    with zipfile.ZipFile(docx_path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/"):
                continue
            target = images_dir / Path(name).name
            target.write_bytes(archive.read(name))
            images.append({"source": name, "path": str(target)})

    return images


def markdown_from_blocks(blocks: list[dict]) -> str:
    lines = []
    for index, block in enumerate(blocks, start=1):
        if block["type"] == "paragraph":
            lines.append(f"<!-- block {index}: paragraph -->")
            lines.append(block["text"])
            lines.append("")
        elif block["type"] == "table":
            lines.append(f"<!-- block {index}: table -->")
            rows = block["rows"]
            if rows:
                width = max(len(row) for row in rows)
                normalized = [row + [""] * (width - len(row)) for row in rows]
                header = normalized[0]
                lines.append("| " + " | ".join(cell.replace("\n", "<br>") for cell in header) + " |")
                lines.append("| " + " | ".join("---" for _ in header) + " |")
                for row in normalized[1:]:
                    lines.append("| " + " | ".join(cell.replace("\n", "<br>") for cell in row) + " |")
            lines.append("")
    return "\n".join(lines)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: extract_docx.py <source.docx> <output_dir>", file=sys.stderr)
        return 2

    source = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])
    output_dir.mkdir(parents=True, exist_ok=True)

    local_docx = output_dir / source.name
    if source.resolve() != local_docx.resolve():
        shutil.copy2(source, local_docx)

    document = Document(str(local_docx))
    blocks = list(iter_block_items(document))
    images = extract_images(local_docx, output_dir)
    payload = {
        "source": str(source),
        "local_docx": str(local_docx),
        "block_count": len(blocks),
        "image_count": len(images),
        "blocks": blocks,
        "images": images,
    }

    (output_dir / "docx-extraction.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    (output_dir / "docx-extraction.md").write_text(
        markdown_from_blocks(blocks),
        encoding="utf-8",
    )
    print(json.dumps({"blocks": len(blocks), "images": len(images)}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())


